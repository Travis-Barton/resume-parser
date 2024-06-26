import openai
import streamlit as st
from utils import file_reader
from llm_class import LLMUtils
from openai import ChatCompletion
from docx import Document
from docx.shared import Pt
import os
from langchain.vectorstores import FAISS
from langchain.embeddings.openai import OpenAIEmbeddings
import dotenv
from docxcompose.composer import Composer
from docx.oxml.ns import qn
import tiktoken

tokenizer = tiktoken.get_encoding("cl100k_base")

fais_db = "vectorstore2"

replace_dict = {
    "summary": "SUMMARY_REPLACE",
    "skills_and_tech": "SKILL_REPLACE",
    "professional_experience": "JOBEXPERIENCE_REPLACE",
    "education": "EDUCATION_REPLACE",
    "certifications": "CERTIFICATION_REPLACE",
    "awards": "AWARD_REPLACE"
}

HEADER_DICT = {
    "summary": "SUMMARY",
    "skills_and_tech": "SKILLS AND TECHNOLOGIES",
    "professional_experience": "PROFESSIONAL EXPERIENCE",
    "education": "EDUCATION",
    "certifications": "CERTIFICATIONS",
    "awards": "AWARDS"
}


def list_number(doc, par, prev=None, level=None, num=False):
    def style_xpath(prefer_single=True):
        style = par.style.style_id
        return ('w:abstractNum[{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]]/@w:abstractNumId').format(
            style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        type = 'bullet'
        return ('w:abstractNum[{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]]/@w:abstractNumId').format(
            type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    xpath_options = {
        True: {
            'single': 'count(w:lvl)=1 and ',
            'level': 0},
        False: {
            'single': '',
            'level': level},
    }

    if prev is None or not prev._p.pPr or not prev._p.pPr.numPr or not prev._p.pPr.numPr.numId:
        level = 0 if level is None else level
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        anum = get_abstract_id()
        num = numbering.add_num(anum)
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        num = num.numId
    else:
        level = prev._p.pPr.numPr.ilvl.val if level is None else level
        num = prev._p.pPr.numPr.numId.val

    if num:
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level


def add_list(document, items):
    paragraphs = []
    for item in items:
        if isinstance(item, str):
            par = document.add_paragraph(item, style='List Paragraph')
            list_number(document, par, prev=paragraphs[-1] if paragraphs else None)
            paragraphs.append(par)
        elif isinstance(item, list):
            for sub_item in item:
                par = document.add_paragraph(sub_item, style='List Paragraph')
                list_number(document, par, prev=paragraphs[-1], level=1)
                paragraphs.append(par)
    return paragraphs


def get_binary_file_downloader_html(bin_file, file_label='File'):
    import base64
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}">{file_label}</a>'
    return href


def get_binary_file_downloader(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return data


def format_professional_experience(content, doc):
    content = content.strip()
    lines = content.split('\n\n')
    bold_lines = []
    bullet_list = []
    for brick in lines:
        for idx, line in enumerate(brick.split('\n')):
            if idx < 2:
                doc.add_paragraph(line, style='Travis Normal Bold 2')
            elif line == '\n':
                doc.add_paragraph(line)
            else:
                doc.add_paragraph(line, style='List Bullet 2')


def format_skills_and_tech(content, doc):
    content.replace('\n\n', '\n')
    for line in content.split('\n'):
        doc.add_paragraph(line, style='List Bullet 2')
    # for paragraph in doc.paragraphs:
    #     replace_text_in_paragraph(paragraph, replace_dict['skills_and_tech'], content.strip())


def format_education(content, doc):
    doc.add_paragraph(content, style='Travis Normal')


def format_summary(content, doc):
    doc.add_paragraph(content, style='Travis Normal')


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


def get_binary_file_downloader_html(bin_file, file_label='File'):
    import base64
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}">{file_label}</a>'
    return href


def create_doc(doc_path, sections, section_results, model='fine_tuned'):
    if model == 'fine_tuned':
        fine_tuned_parser(section_results, doc_path)
    elif model == 'gpt_16k':
        gpt_16k_parser(section_results, doc_path)


def content_markup(content, section):
    # content = content.replace("\n- ", "\n• ")
    # content = content.replace("\t- ", "\t• ")
    content = content.replace("\t- ", "\t")
    content = content.replace("\n- ", "\n ")
    content = content.replace("\nSUMMARY\n", "")
    content = content.replace("\nSKILLS AND TECHNOLOGY\n", "")
    content = content.replace("\nPROFESSIONAL EXPERIENCE\n", "")
    content = content.replace("\nEDUCATION\n", "")
    content = content.replace('```education', '')
    content = content.replace("`", "").replace('skills_and_tech', '')
    content = content.replace('professional_experience', '')
    if content.startswith(':\n'):
        content = content[2:]
    if content.startswith('education'):
        content = content[9:]
    if section == 'professional_experience':
        content = content.split('\nEducation')[0]
        content = content.split('\nSKILLS\n')[0]
        content = content.split('\nSkills Summary\n')[0]
    return content


def fine_tuned_parser(section_results, doc_path):
    doc_dict = {
        'summary': Document('templates/SUMMARY.docx'),
        'skills_and_tech': Document('templates/SKILLS AND TECHNOLOGIES.docx'),
        'professional_experience': Document('templates/PROFESSIONAL EXPERIENCE.docx'),
        'education': Document('templates/EDUCATION.docx'),
    }
    master_doc = Document('templates/AIM Profile - Template.docx')
    composer = Composer(master_doc)
    for section, content in section_results.items():
        doc = doc_dict[section]

        content = content_markup(content, section)
        if section == 'professional_experience':
            content = content.replace('•', '')
            format_professional_experience(content, doc)
            continue
        elif section == 'skills_and_tech':
            format_skills_and_tech(content, doc)
            continue
        elif section == 'education':
            format_education(content, doc)
            continue
        elif section == 'summary':
            format_summary(content, doc)
            continue

        # for paragraph in doc.paragraphs:
        #     replace_text_in_paragraph(paragraph, replace_dict[section], content.strip())

    for key, value in doc_dict.items():
        # merge the docs
        composer.append(value)
    # save the doc
    composer.save(doc_path)
    final_doc = Document(doc_path)
    for paragraph in final_doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
        paragraph.space_before = 0  # Set before spacing to 0
        paragraph.space_after = 0  # Set after spacing to 0

        # For tables
    for table in final_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                    paragraph.space_before = 0  # Set before spacing to 0
                    paragraph.space_after = 0  # Set after spacing to 0

        # For headers and footers
    for section in final_doc.sections:
        # Header
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri'
            paragraph.space_before = 0  # Set before spacing to 0
            paragraph.space_after = 0  # Set after spacing to 0
        # Footer
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri'
            paragraph.space_before = 0  # Set before spacing to 0
            paragraph.space_after = 0  # Set after spacing to 0
    final_doc.save(doc_path)


def gpt_16k_parser(section_results, doc_path):
    doc = Document('templates/AIM Profile - Template.docx')
    for section, content in section_results.items():
        content = content.replace('•', '')
        content = content_markup(content, section)
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, replace_dict[section], content.strip())
        # doc.add_paragraph(content)
        # style = doc.styles['Normal']
    doc.save(doc_path)


# Sidebar for model selection
with st.sidebar:
    model_select = st.selectbox('Select the right model', ['fine-tuned-gpt-3.5-turbo-4k',
                                                           'fine_tuned_gpt-3.5-turbo-16k-no-dot',
                                                           'gpt-3.5-turbo-16k', ])
    if model_select == 'gpt-3.5-turbo-16k':
        st.warning('Using the fine-tuned model is recommended for better results.')
    api_key = st.text_input('Enter your OpenAI API key', type='password')
    if api_key:
        openai.api_key = api_key
        embeddings = OpenAIEmbeddings()
        llm_util = LLMUtils(model=ChatCompletion(),
                            embedding_source=FAISS.load_local(fais_db, embeddings=embeddings),
                            model_names=model_select)
    if model_select == 'fine-tuned-gpt-3.5-turbo-4k':
        openai.api_key = st.secrets['OPENAI_API_KEY']

# Main content
st.title("AIM Resume Processor")

# Upload resume
check = st.checkbox('Paste Mode', value=False)
if not check:
    resume_inputer = st.file_uploader('Upload the resume here', type=['doc', 'docx', 'pdf', 'txt'])
else:
    resume_inputer = st.text_area('Paste the resume here')

if api_key:
    st.success('Press R to rerun the analysis')
    if resume_inputer:
        try:
            resume = file_reader(resume_inputer) if not check else resume_inputer
            print(resume)
            if (not resume) or (resume == '') or (len(resume) < 100):
                st.error("The resume returned empty. Please try again.")
                st.stop()
        except Exception as e:
            st.error(f'Error reading file: {e}')
            st.stop()
        # sections = ['summary', 'skills_and_tech', 'professional_experience', 'education', 'certifications', 'awards']
        sections = ['summary', 'skills_and_tech', 'professional_experience', 'education']
        section_results = {}
        resume_len = len(tokenizer.encode(resume))
        if resume_len > 4096 and model_select == 'fine-tuned-gpt-3.5-turbo-4k':
            st.error(
                'Resume is too long. Please upload a shorter resume or use the GPT-3.5-turbo-16k model (see sidebar)')
            st.stop()

        for section in sections:
            result = llm_util.extract_section(resume=resume, section=section)
            section_results[section] = result
            with st.expander(f"Extracted {section.replace('_', ' ').upper()}",
                             expanded=False):
                st.markdown(f"# {section.replace('_', ' ').upper()}\n {result}",
                            unsafe_allow_html=True)  # Display extracted content
        # Create a DOCX from the extracted sections
        # read doucment
        # doc = Document('AIM Profile template.docx')
        doc_path = 'resume.docx'
        # model = 'fine_tuned' if model_select == 'fine-tuned-gpt-3.5-turbo-4k' else 'gpt_16k'
        model = 'fine_tuned'
        create_doc(doc_path=doc_path, sections=sections, section_results=section_results, model=model)

        with st.sidebar:
            data = get_binary_file_downloader(doc_path)
            st.download_button(label='Download DOCX', data=data, file_name='consolidated.docx',
                               mime='application/octet-stream')

        data = get_binary_file_downloader(doc_path)
        st.download_button(label='Download DOCX', data=data, file_name='consolidated.docx',
                           mime='application/octet-stream', key='download_button2')
else:
    st.markdown('__Please enter your OpenAI API key in the sidebar__')

with st.sidebar:
    st.markdown('---')
    st.markdown('Whats new?\n1. A new fine-tuned model for not including periods.\n2. a better looking download button.')
    st.markdown(
        "<br><br><br><br>Something broken? <br>[File an issue](https://github.com/Travis-Barton/resume-parser/issues) or "
        "reach out to me <a href='mailto:me@travisbarton.com?subject=Resume Parser'>by email</a>",
        unsafe_allow_html=True)
    st.markdown('<br>_Made by [Travis Barton Consulting](https://www.travisbarton.com)_',
                unsafe_allow_html=True)
